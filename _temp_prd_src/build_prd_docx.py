# -*- coding: utf-8 -*-
"""Generate reorganized PRD Word document from extracted requirements."""

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_meta_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("游戏账号交易平台\n产品需求说明书（PRD）")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    add_heading(doc, "文档信息", 1)
    add_meta_table(
        doc,
        [
            ("产品名称", "游戏账号交易平台"),
            ("文档类型", "产品需求说明书"),
            ("文档版本", "V1.1"),
            ("编写日期", date.today().isoformat()),
            ("需求来源", "《济源刘总游戏账号交易平台（4端）_PRD与报价整合版》整理稿"),
            ("适用端口", "PC Web、H5、Android APP、iOS APP、管理后台"),
            ("说明", "本版为需求归并稿，费用与商务条款不在此文件展开（另附报价/合同）。"),
        ],
    )

    doc.add_paragraph()
    add_heading(doc, "修订记录", 2)
    rt = doc.add_table(rows=2, cols=5)
    rt.style = "Table Grid"
    hdr = ["版本", "日期", "作者", "变更说明", "影响范围"]
    for j, h in enumerate(hdr):
        rt.rows[0].cells[j].text = h
    rt.rows[1].cells[0].text = "V1.1"
    rt.rows[1].cells[1].text = date.today().isoformat()
    rt.rows[1].cells[2].text = "项目组"
    rt.rows[1].cells[3].text = "按原稿重排章节、去口语化表述、补齐范围/验收/待定项"
    rt.rows[1].cells[4].text = "全文"

    doc.add_paragraph()
    add_heading(doc, "1 背景与目标", 1)
    add_heading(doc, "1.1 背景", 2)
    add_para(
        doc,
        "建设面向游戏账号、游戏金币、游戏道具及中介担保的综合交易平台，面向买家、卖家、代理、客服、运营、财务及后台管理员，覆盖商品发布与审核、浏览与筛选、下单支付、订单协同（群聊）、换绑验号、售后退款、提现分账、推广返利等闭环流程。",
    )
    add_para(
        doc,
        "甲方对多端界面呈现要求较高：PC、H5、Android、iOS 需提供高清素材适配与一致的暗黑沉浸式视觉语言（具体以签字确认的设计稿为准）。移动端不得仅以 PC 页面等比缩放替代独立布局。",
    )

    add_heading(doc, "1.2 业务目标", 2)
    goals = [
        "统一承载账号、金币、道具、中介担保等交易类型；",
        "买家可按游戏、区服、价格区间及账号属性高效检索并下单；",
        "卖家可自助发布商品，经平台审核后上架展示；",
        "支持担保交易、包赔套餐、电子合同、订单群聊与人工售后；",
        "支持主流支付方式、退款、四方支付分账及卖家提现；",
        "支持邀请推广、优惠券、积分商城、卡包、收藏、足迹、出价等运营能力；",
        "支持代理授权与人工估号，支撑回收、代理回收及客服介入场景。",
    ]
    for i, g in enumerate(goals, 1):
        doc.add_paragraph(f"{i}. {g}", style="List Number")

    add_heading(doc, "1.3 产品目标", 2)
    add_para(
        doc,
        "在合规与安全可控前提下，形成高辨识度的游戏风视觉体系，并以实名、隐私声明、担保与合同等手段降低交易不确定性；后台通过统一属性模板与审核流程保证商品配置效率与一致性。",
    )

    add_heading(doc, "1.4 范围说明", 2)
    scope = doc.add_table(rows=3, cols=2)
    scope.style = "Table Grid"
    scope.rows[0].cells[0].text = "本期范围（In）"
    scope.rows[0].cells[1].text = (
        "四端用户前台 + 管理后台；商品与订单全流程；支付/退款/分账/提现（以实际支付资质与通道为准）；"
        "IM 群聊与私信；实名与人脸（以选型的合规服务商能力为准）；电子签（法大大或同类，待法务确认）；"
        "运营配置（游戏/属性/优惠券/推广等）。"
    )
    scope.rows[1].cells[0].text = "非本期范围（Out）"
    scope.rows[1].cells[1].text = (
        "复杂竞价大厅、自动估价与智能推荐、账号安全自动化检测、直播式展示、多级代理结算、"
        "强 AI 客服、多语言与海外支付、自研风控模型等；若后续纳入，需单独立项评估。"
    )
    scope.rows[2].cells[0].text = "假设与约束"
    scope.rows[2].cells[1].text = (
        "第三方支付、IM、实名、人脸、电子签等以外部服务商能力与甲方资质为前提；"
        "虚拟商品交易的法务定性、合同模板及费率规则须甲方书面确认。"
    )

    doc.add_paragraph()
    add_heading(doc, "2 术语与定义", 1)
    term_rows = [
        ("商品", "待交易的账号、金币、道具或服务条目，含属性、图片及审核状态。"),
        ("订单", "买卖双方围绕单次交易生成的履约单据，关联支付、交付与售后状态。"),
        ("担保交易", "平台提供履约监督与争议处理机制的交易形态（细则见协议文本）。"),
        ("包赔/保险套餐", "订单可选增值服务类型，赔付规则以后台配置及前台展示为准。"),
        ("出价/砍价", "买家出价、卖家响应的议价流程，可涉及定金规则（比例可配置）。"),
        ("订单群聊", "基于订单生成的即时沟通频道，用于交付协同。"),
        ("分账", "支付机构规则下的多方资金拆分，具体角色与比例待财务确认。"),
        ("代理", "经后台授予的身份，可在授权范围内发布与报价；毁约次数与处罚规则可配置。"),
        ("人工估号", "用户提交资料由客服评估回收报价的流程。"),
    ]
    terms = doc.add_table(rows=1 + len(term_rows), cols=2)
    terms.style = "Table Grid"
    terms.rows[0].cells[0].text = "术语"
    terms.rows[0].cells[1].text = "说明"
    for i, (k, v) in enumerate(term_rows, start=1):
        terms.rows[i].cells[0].text = k
        terms.rows[i].cells[1].text = v

    doc.add_paragraph()
    add_heading(doc, "3 角色与权限", 1)
    role_data = [
        ("游客", "浏览公开内容", "首页、列表、帮助与安全说明"),
        ("注册用户", "账户建立与资料维护", "登录、浏览、收藏、发起实名"),
        ("买家", "选购与履约", "下单、支付、合同签署、售后申请"),
        ("卖家", "发布与交付", "发布、上下架、履约、提现申请"),
        ("代理", "授权范围内发布/报价", "代理商品、报价、毁约记录受规则约束"),
        ("客服", "协同与售后", "估号受理、群聊协同、工单处理"),
        ("审核员", "内容/资金审核", "商品、提现、退款、实名材料审核"),
        ("财务", "资金核对", "支付流水、分账、对账视图"),
        ("运营", "配置与活动", "游戏、属性、优惠券、推广配置"),
        ("系统管理员", "权限与参数", "账号角色、菜单、系统参数、日志"),
    ]
    roles = doc.add_table(rows=1 + len(role_data), cols=3)
    roles.style = "Table Grid"
    roles.rows[0].cells[0].text = "角色"
    roles.rows[0].cells[1].text = "职责摘要"
    roles.rows[0].cells[2].text = "核心权限（示例）"
    for i, row in enumerate(role_data, start=1):
        for j, val in enumerate(row):
            roles.rows[i].cells[j].text = val

    doc.add_paragraph()
    add_heading(doc, "4 业务流程", 1)

    flows = [
        (
            "4.1 注册登录",
            [
                "用户进入登录/注册入口，录入手机号并获取验证码；",
                "勾选《用户协议》《隐私政策》，未勾选不允许继续；",
                "校验通过后创建账户并登录；",
                "首次涉及交易前完成实名与人脸（规则待风控确认）；",
                "验证码具备发送频控、有效期与失败锁定策略。",
            ],
        ),
        (
            "4.2 买家购买（账号/金币/道具/中介担保）",
            [
                "选择交易类型与游戏（支持检索、热门、历史选择）；",
                "在列表页筛选（游戏、区服、端口、价格、等级等，以游戏配置为准）；",
                "进入详情核对属性、图片、是否支持合同等；",
                "可收藏后下单；填写/确认联系方式与游戏内信息；",
                "选择包赔套餐、优惠券、推荐码与支付方式；",
                "勾选交易协议并完成支付；",
                "生成订单群聊，双方完成发货、换绑、验号与确认收货；",
                "异常进入退款/售后或人工客服通道。",
            ],
        ),
        (
            "4.3 卖家发布",
            [
                "选择类型与游戏，填写文案、亮点、截图及动态属性字段；",
                "上传图片（命名与预览规则按设计稿）；系统生成商品编号；",
                "预览详情页后提交审核；后台通过后上架展示；",
                "卖家可在“我的卖出”查看出售中、已售出、已下架等状态。",
            ],
        ),
        (
            "4.4 议价（出价）",
            [
                "买家发起报价；平台可按规则冻结定金（默认参照示例比例，后台可配置）；",
                "双方私信沟通，支持图片/文件；",
                "卖家同意后买方支付尾款进入正式订单流程；",
                "拒绝或超时场景下的定金处理规则由后台配置并在前台明示。",
            ],
        ),
        (
            "4.5 合同签署",
            [
                "订单进入交付阶段后生成合同草案；",
                "对接电子签服务（如法大大），双方完成签署；",
                "卖家提现前须完成合同约定动作（如有）。",
            ],
        ),
        (
            "4.6 提现",
            [
                "卖家在钱包查看余额、可提现额度与手续费（如开启）；",
                "绑定提现账户，提交申请；后台审核通过打款并留痕。",
            ],
        ),
        (
            "4.7 人工估号",
            [
                "用户选择游戏并提交属性与截图；",
                "客服评估并反馈回收报价；用户确认后转入私聊或订单群协同。",
            ],
        ),
        (
            "4.8 代理授权",
            [
                "后台授予代理身份；代理在授权品类下发布“平台交易/代理回收”类商品；",
                "卖家接受报价后拉群交易；毁约计数与处罚（如超过阈值取消代理身份、扣减比例）按配置执行。",
            ],
        ),
    ]

    for heading, steps in flows:
        add_heading(doc, heading, 2)
        for step in steps:
            doc.add_paragraph(step, style="List Bullet")

    doc.add_paragraph()
    add_heading(doc, "5 功能需求（按端口）", 1)
    add_para(
        doc,
        "以下按端口汇总能力。细到字段级的配置以《后台游戏/属性模板》及签字设计稿为准；未列出的交互细节在实施阶段以原型评审纪要为准。",
    )

    pf = doc.add_table(rows=6, cols=2)
    pf.style = "Table Grid"
    pf.rows[0].cells[0].text = "端口"
    pf.rows[0].cells[1].text = "功能范围"
    port_rows = [
        (
            "PC Web",
            "首页、游戏与商品浏览、筛选排序、详情、下单支付、订单、消息、个人中心、推广入口等桌面端完整能力。",
        ),
        (
            "H5",
            "面向移动浏览器/分享落地：注册登录、浏览、下单支付、个人中心、推广邀请等；独立移动信息架构与导航。",
        ),
        (
            "Android APP",
            "交易主流程、推送、钱包提现、订单群聊、实名与人脸、高清资源适配。",
        ),
        (
            "iOS APP",
            "与安卓对等的核心交易能力；适配安全区与高分辨率屏幕。",
        ),
        (
            "管理后台",
            "用户、游戏与属性模板、商品审核、订单/支付/退款、提现审核、优惠券与积分商品、推广奖励、代理与权限、内容与日志。",
        ),
    ]
    for i, (a, b) in enumerate(port_rows, start=1):
        pf.rows[i].cells[0].text = a
        pf.rows[i].cells[1].text = b

    modules = [
        "登录注册：手机号验证码、协议勾选、会话与安全策略。",
        "首页：买卖入口、平台安全说明、帮助、回收/积分商城入口等（以稿为准）。",
        "买点什么：类型分流、游戏选择、列表、详情、收藏。",
        "卖点什么：类型选择、动态表单、图片上传、预览、提交审核。",
        "下单与支付：保险套餐、优惠券、推荐码、协议勾选、支付结果页。",
        "订单中心：买家/卖家视图、状态筛选、详情、售后入口。",
        "消息中心：系统通知、订单群聊、私信、出价卡片、附件传输。",
        "我的：资料、实名、钱包、提现记录、卡包、收藏、足迹、出价记录、交易记录。",
        "推广邀请：邀请码/链接、关系绑定、奖励规则展示。",
        "支付板块：微信支付、支付宝、退款、四方渠道与分账（以实际开通为准）。",
        "后台配置：游戏 Logo、筛选字段、商品属性模板、优惠券、积分商品、风控与审核流。",
    ]
    add_heading(doc, "5.1 用户端共性模块", 2)
    for m in modules:
        doc.add_paragraph(m, style="List Bullet")

    doc.add_paragraph()
    add_heading(doc, "6 UI/UX 与视觉交付要求", 1)
    add_para(
        doc,
        "总体风格：暗黑沉浸式、游戏场景化背景、玻璃拟态导航与卡片、紫色/青色强调色（可随品牌微调）。PC 与移动端共享视觉语言但需分别出稿。",
    )
    ux_req = [
        "高分辨率素材分端提供，禁止简单拉伸导致糊边与形变；",
        "列表/详情大图采用渐进式加载与占位策略；",
        "动效应轻量并具备降级策略，核心支付/实名流程避免过度动效；",
        "不建议以长时视频作为全站常态背景（性能与耗电）；",
        "关键页面需走查大屏/普屏/移动高分屏可读性（价格、状态、按钮对比度）；",
        "设计交付须包含：组件库规范、切图标注、交互说明、空状态与错误态。",
    ]
    for u in ux_req:
        doc.add_paragraph(u, style="List Bullet")

    add_heading(doc, "6.1 美术与 3D 资源（如需）", 2)
    add_para(
        doc,
        "重要首屏与高转化页面可采用 3D 静帧或预渲染图 + 前端轻动效的组合，以平衡观感与性能。资源来源（客户素材、正版商用、原创、外协）及版权责任由甲方确认。",
    )

    doc.add_paragraph()
    add_heading(doc, "7 数据与状态", 1)
    add_heading(doc, "7.1 核心对象", 2)
    objs = "用户、游戏、商品、订单、钱包流水、优惠券、推广奖励、消息/会话。"
    add_para(doc, objs)
    add_heading(doc, "7.2 状态机（摘要）", 2)
    sm = doc.add_table(rows=5, cols=2)
    sm.style = "Table Grid"
    sm_data = [
        ("商品状态", "草稿、待审核、上架、下架、成交关闭等（以后台字典为准）"),
        ("订单状态", "待支付、待发货/待交付、履约中、待确认、完成、关闭、售后中等"),
        ("提现状态", "待审核、处理中、成功、失败、驳回"),
        ("售后状态", "申请中、协商中、退款中、完结/关闭"),
    ]
    sm.rows[0].cells[0].text = "领域"
    sm.rows[0].cells[1].text = "说明"
    for i, (a, b) in enumerate(sm_data, start=1):
        sm.rows[i].cells[0].text = a
        sm.rows[i].cells[1].text = b

    doc.add_paragraph()
    add_heading(doc, "8 非功能需求", 1)
    nfr_data = [
        ("性能", "首页首屏关键信息建议 3s 内可见；列表分页/虚拟滚动；图片多级尺寸与缓存。"),
        ("安全", "全站 HTTPS；验证码防刷；支付/退款/分账回调验签；资金与审核操作留痕；敏感信息脱敏；文件上传校验。"),
        ("合规", "协议与风险提示完备；实名与人脸授权；电子合同主体与模板经法务确认；资金方案符合支付机构规则。"),
        ("兼容", "覆盖主流浏览器与系统版本矩阵（立项时冻结清单）。"),
        ("可运维", "日志、审计、告警与后台操作追溯；关键配置可视化。"),
    ]
    nfr = doc.add_table(rows=1 + len(nfr_data), cols=2)
    nfr.style = "Table Grid"
    nfr.rows[0].cells[0].text = "类别"
    nfr.rows[0].cells[1].text = "要求"
    for i, (a, b) in enumerate(nfr_data, start=1):
        nfr.rows[i].cells[0].text = a
        nfr.rows[i].cells[1].text = b

    doc.add_paragraph()
    add_heading(doc, "9 集成与第三方", 1)
    add_para(
        doc,
        "IM、支付、实名/人脸、电子签、对象存储与 CDN、短信等均通过合规供应商接入。具体厂商与账号由甲方准备；联调环境需双方共同申请。接口错误码、重试与对账策略在《接口说明》中单列。",
    )

    doc.add_paragraph()
    add_heading(doc, "10 验收标准（摘要）", 1)
    acc = doc.add_table(rows=4, cols=2)
    acc.style = "Table Grid"
    acc.rows[0].cells[0].text = "类别"
    acc.rows[0].cells[1].text = "要点"
    acc_data = [
        (
            "UI/前端",
            "四端核心页面还原签字设计稿；暗黑风一致；高分屏清晰度达标；核心流程无阻断性布局错误。",
        ),
        (
            "业务功能",
            "注册登录带协议校验；买卖发布与审核闭环；下单支付（至少一种测试通道联调）；订单与群聊联动；钱包与提现审核；推广与优惠券主路径可用。",
        ),
        (
            "后台",
            "RBAC 生效；游戏/属性/优惠券可配置；审核与资金操作可追溯；报表/列表分页与导出（如约定）可用。",
        ),
    ]
    for i, (a, b) in enumerate(acc_data, start=1):
        acc.rows[i].cells[0].text = a
        acc.rows[i].cells[1].text = b

    doc.add_paragraph()
    add_heading(doc, "11 里程碑建议", 2)
    add_para(doc, "可与报价拆解对齐，以下为常用阶段划分：需求冻结与交互评审 → UI 视觉与素材 → 前端框架与页面还原 → 后端与后台 → 联调与安全测试 → 试运行与验收上线 → 维保交接。")

    doc.add_paragraph()
    add_heading(doc, "12 待定问题清单", 1)
    open_q = [
        "四方支付具体服务商与费率、分账角色清单及对账接口口径；",
        "IM 选型（腾讯云 IM 或其他）与消息留存周期；",
        "电子签合同模板字段与签署顺序；",
        "包赔套餐的法律文案与赔付上限；",
        "代理毁约阈值与违约金比例的最终数值；",
        "人脸识别服务商及活体策略；",
        "是否需要跨境或多币种（默认本期不做）；",
        "管理后台是否需要深色主题强制一致或允许浅色高效主题。",
    ]
    for q in open_q:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_paragraph()
    add_heading(doc, "附录 A 技术路线说明（摘录，待甲方确认）", 1)
    add_para(
        doc,
        "原整合稿给出的路线为：PC/H5 侧采用主流 Web 技术栈实现服务端渲染或等价方案；移动端采用跨平台框架以降低双端分叉成本；后台采用 Web 管理端；后端采用 Java；数据层采用关系型数据库与缓存，并结合对象存储与 CDN。"
        "最终选型以实现签字 PRD、性能与安全指标为前提，在项目启动评审会上确认。",
    )

    out = "/Users/liyan/Desktop/4月份项目/我的 AI 项目/济源游戏账号交易平台_PRD/游戏账号交易平台_PRD_v1.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
