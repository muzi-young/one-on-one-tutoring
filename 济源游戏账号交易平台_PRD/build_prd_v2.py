# -*- coding: utf-8 -*-
"""从整合稿提取结果生成 PRD v2（产品经理书面表述，表格信息量保留）。"""
from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

JSON_PATH = Path(__file__).resolve().parent.parent / "_temp_prd_src" / "_extracted.json"
OUT_PATH = Path(__file__).resolve().parent / "游戏账号交易平台_PRD_v2.docx"


def add_title(doc: Document, text: str, pt: int = 18) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(pt)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table_grid(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell_text = row[j] if j < len(row) else ""
            t.rows[i].cells[j].text = (cell_text or "").strip()


def strip_bullet(s: str) -> str:
    s = s.strip()
    return re.sub(r"^[-\d\.]+\s*", "", s).strip() if s else s


def load_tables() -> list[list[list[str]]]:
    with open(JSON_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data["tables"]


def doc_meta_rows(T: list) -> list[list[str]]:
    base = [list(r) for r in T[0]]
    for row in base:
        if row[0] == "文档版本":
            row[1] = "V2.0"
        if row[0] == "项目阶段":
            row[1] = "需求确认、UI 设计、技术评估、开发排期（商务测算见独立附件，不以此表为报价依据）"
        if row[0] == "覆盖端口":
            row[1] = (
                "用户端四端：PC Web、H5、Android APP、iOS APP；运营侧：管理后台（Web）。"
                "说明：业务上常称“四端”指用户触点；本需求同时包含后台，共五套前端交付面。"
            )
        if row[0] == "重点要求":
            row[1] = (
                "高保真 UI、高清素材分端交付、暗黑沉浸式游戏风视觉语言表达；"
                "交易/支付/售后/担保链路可追溯。"
            )
    base.append(["需求溯源", "《游戏账号交易平台定制开发 PRD 文档》整合版（内部归并稿，含原报价卷宗索引字段已剔除金额）"])
    base.append(["适用范围", "立项评审、交互/视觉评审、研发排期、测试验收（不含商务合同条款）"])
    return base


def revision_table() -> list[list[str]]:
    return [
        ["版本", "日期", "作者", "变更说明", "影响范围"],
        [
            "V1.0",
            "—",
            "项目组",
            "整合稿基线（来源 Word）",
            "全文",
        ],
        [
            "V2.0",
            date.today().isoformat(),
            "项目组",
            "保留整合稿表格与流程信息量；改写为可评审书面语；API 改为职责摘要（不含路径）；正文不收口径价",
            "全文",
        ],
    ]


def api_duty_table(T: list, domains: list[str], start: int, end: int) -> list[list[str]]:
    rows: list[list[str]] = [["业务域", "职责要点（实施阶段锁定接口形态，本文不承载 URL）"]]
    for name, idx in zip(domains, range(start, end)):
        tbl = T[idx]
        for r in tbl[1:]:
            if len(r) < 3:
                continue
            duty = (r[2] or "").strip()
            verb = (r[0] or "").strip()
            note = f"{duty}" if not verb else f"【{verb}】{duty}"
            rows.append([name, note])
    return rows


def ascii_block_lines() -> str:
    """来源整合稿中的架构示意 ASCII，原样收录便于追溯。"""
    return "\n".join(
        [
            "用户端 PC Web / H5",
            "│",
            "├── React + Next.js + TypeScript",
            "│",
            "Android / iOS APP",
            "│",
            "├── Flutter 3 + Dart",
            "│",
            "管理后台",
            "│",
            "├── Vue 3 + TypeScript + Ant Design Vue",
            "│",
            "▼",
            "统一 API 网关 / 后端接口层",
            "│",
            "├── 用户与认证服务",
            "├── 游戏与商品服务",
            "├── 订单交易服务",
            "├── 支付退款与分账服务",
            "├── 钱包提现服务",
            "├── 优惠券与推广服务",
            "├── IM 群聊对接服务",
            "├── 实名认证与电子签服务",
            "└── 后台权限与审计服务",
            "│",
            "▼",
            "MySQL 8 + Redis + OSS/COS + CDN + 第三方服务",
        ]
    )


def main() -> None:
    T = load_tables()
    with open(JSON_PATH, "r", encoding="utf-8") as f:
        pmap = {x["i"]: x["text"] for x in json.load(f)["paras"]}
    doc = Document()

    add_title(doc, "游戏账号交易平台\n产品需求说明书（PRD）\nV2.0")

    h(doc, "文档信息", 1)
    table_grid(doc, doc_meta_rows(T))

    h(doc, "修订记录", 2)
    table_grid(doc, revision_table())

    # —— 背景 ——
    h(doc, "1 背景", 1)
    p(
        doc,
        "拟建设面向游戏账号、游戏金币、游戏道具及中介担保的综合交易平台，服务对象包含买家、卖家、代理、客服、"
        "运营、财务及后台管理员。系统需支撑商品发布、审核、浏览检索、下单支付、订单协同（群聊）、换绑验号、"
        "售后退款、提现分账、推广返利等闭环能力。",
    )
    p(
        doc,
        "甲方对界面呈现与分辨率表现有明确验收预期：PC、H5、Android、iOS 需提供高清素材与独立布局（非单一"
        "分辨率拉伸替代设计）。对标素材为评审基准：分辨率 3360×1738、时长约 166.65s 的参考视频，用于共识"
        "“暗黑游戏场景 + 玻璃拟态导航 + 高对比订单侧栏 + 分步交易表单”的目标示意；最终以签字设计稿为准。",
    )

    # —— 目标 ——
    h(doc, "2 目标", 1)
    h(doc, "2.1 业务目标", 2)
    bullets(
        doc,
        [
            "统一承载账号、金币、道具、中介担保等交易类型。",
            "买家可按游戏、区服、价格区间及账号属性等条件筛选与排序。",
            "卖家可自助发布商品，经平台审核后上架展示。",
            "支持担保交易、包赔套餐、电子合同、订单群聊与人工售后。",
            "支持微信支付、支付宝、退款、四方支付分账及卖家提现。",
            "支持推广邀请、优惠券、积分商城、卡包、收藏、足迹、出价等运营能力。",
            "支持代理授权与人工估号，覆盖平台回收、代理回收及客服介入场景。",
        ],
    )
    h(doc, "2.2 产品目标", 2)
    bullets(
        doc,
        [
            "在合规约束内形成可识别的暗黑游戏风品牌界面，并保证关键交易信息的可读性与对比度。",
            "通过协议提示、实名/人脸、担保与合同等机制降低履约不确定性（规则以法务文本为准）。",
            "通过统一游戏配置、属性模板与审核流降低商品配置成本并保持前后台口径一致。",
            "通过 IM、订单状态与售后入口提升履约过程的可观察性。",
        ],
    )

    # —— 范围 ——
    h(doc, "3 端口与范围", 1)
    h(doc, "3.1 端口划分", 2)
    p(
        doc,
        "“四端”在业务语境中多指用户触点（PC Web、H5、Android、iOS）。本文件同时包含管理后台；"
        "工作量评估与交付验收需按五套前端面分别单列，不得默认后台与前台共用一套交互稿。",
    )
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[1]])

    h(doc, "3.2 本期与二期边界", 2)
    p(doc, "以下为整合稿中对二期能力的边界声明，立项时如需纳入须变更评审并调整基线。")
    bullets(
        doc,
        [
            "更复杂的竞价交易大厅。",
            "自动估价算法与历史成交价智能推荐。",
            "游戏账号安全检测自动化。",
            "直播带货式账号展示。",
            "多级代理体系与代理业绩结算。",
            "AI 客服问答与售后自动分流。",
            "多语言与海外支付渠道。",
            "自研风控模型、黑名单共享库、异常交易识别。",
        ],
    )

    h(doc, "3.3 模块总览（整合稿粒度）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[2]])

    # —— 角色 ——
    h(doc, "4 用户角色与权限矩阵", 1)
    p(
        doc,
        "下表继承整合稿职责划分，实施阶段需在后台 RBAC 中映射为菜单、按钮与数据权限；出现冲突以评审纪要为准。",
    )
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[3]])

    # —— UI 专项 —起
    h(doc, "5 UI／UX 与素材专项", 1)
    p(
        doc,
        "本章单独成章的原因：整合稿将“对标素材的沉浸感 + 高清资源分端 + 动效边界”视为与功能同等重要的验收维度；"
        "评审、排期与测试用例需平行跟踪，避免仅以功能清单通过替代视觉验收。",
    )
    h(doc, "5.1 对标素材元数据（用于效果对齐，非法律承诺）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[4]])
    h(doc, "5.1.1 视觉关键词（整合稿摘录）", 2)
    for i in range(39, 50):
        if i not in pmap:
            continue
        line = pmap[i].strip()
        if not line or line == "核心关键词：" or line == "参考视频解析信息如下：":
            continue
        if line.startswith("-"):
            doc.add_paragraph(strip_bullet(line), style="List Bullet")
        else:
            doc.add_paragraph(line, style="List Bullet")
    h(doc, "5.2 参考画面拆解与落地方式", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[5]])
    h(doc, "5.3 核心页面视觉要求", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[6]])
    h(doc, "5.4 视觉设计规范建议", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[7]])
    h(doc, "5.5 高清素材与切图标准", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[8]])
    h(doc, "5.6 前端高清适配要求（端口）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[9]])
    p(doc, "通用约束（整合稿正文，与上表并用）：")
    bullets(
        doc,
        [strip_bullet(pmap[i]) for i in range(56, 61) if i in pmap and pmap[i].strip().startswith("-")],
    )
    h(doc, "5.7 动效与交互优先级", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[10]])
    p(doc, "动效边界（正文约束，与表内优先级并用）：")
    bullets(
        doc,
        [
            "动效应轻量，不得阻碍支付、下单、上传、聊天等关键路径的可用性。",
            "不建议以长时视频作为常态页面背景（性能、流量与耗电）。",
            "低性能设备必须具备降级策略：关闭粒子/视频层，保留信息与操作控件。",
            "支付、实名、人脸等流程禁止使用干扰确认的过度动效。",
        ],
    )
    h(doc, "5.8 设计与前端交付物", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[11]])
    h(doc, "5.9 移动端独立设计要求", 2)
    p(
        doc,
        "H5、Android、iOS 必须按竖屏信息架构重做主导航与页面层级；禁止以 PC 视觉稿单一缩放作为验收口径。",
    )
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[12]])
    h(doc, "5.10 前端技术栈验收标准（实施前对齐）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[15]])

    # —— 技术路线 ——
    h(doc, "6 拟采用技术路线（待立项评审冻结）", 1)
    p(
        doc,
        "下列内容为整合稿中的主推组合，用语调整为“拟采用”，需在启动会对团队技能、第三方采购与性能指标复核后签字确认。"
        "若变更，应更新修订记录并回灌验收用例。",
    )
    h(doc, "6.1 分端拟采用栈与动机（来源整合稿）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[13]])
    h(doc, "6.2 移动端跨端方案的评估维度（整合稿摘录）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[14]])
    h(doc, "6.3 全链路组件与服务拆分（拟采用）", 2)
    tbl16 = [[c.replace("\n", " ") for c in row] for row in T[16]]
    if tbl16:
        tbl16[0] = ["层级", "拟采用选项（待评审）", "说明"]
    table_grid(doc, tbl16)
    p(doc, "服务拆分逻辑示意（整合稿 ASCII ，便于与下文 API 职责摘要对照）：")
    p(doc, ascii_block_lines())

    h(doc, "6.4 建议人力配置（角色口径，非薪酬报价）", 2)
    p(
        doc,
        "下表为整合稿中对角色投入的建议区间，用于排期与沟通；人员级别与商务条款不在本文展开。",
    )
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[17]])

    # —— 美术 ——
    h(doc, "7 美术资源与交付规格", 1)
    p(
        doc,
        "对标级沉浸页面依赖高质量平面/动效资源；资源版权、授权范围与迭代轮次须在立项附件中由甲方确认。"
        "若资源不足，应在记录风险并下调视觉验收口径，避免无依据承诺。",
    )
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[18]])
    h(doc, "7.1 核心美术资源清单与像素建议", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[19]])
    h(doc, "7.2 来源方式比选", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[20]])
    p(
        doc,
        "组合策略（整合稿实践建议）：甲方授权素材 + 正版商用素材 + 设计师原创 + 生成式辅助探索（须人工统一风格并完成版权审查）。",
    )
    h(doc, "7.3 交付规格", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[21]])
    h(doc, "7.4 验收标准", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[22]])

    # —— 3D ——
    h(doc, "8 3D 与高级视觉（可选方案）", 1)
    p(
        doc,
        "首屏与高转化页可采用 3D 静帧、预渲染序列或局部实时渲染；默认策略建议“预渲染 + 前端轻动效”以平衡观感与性能。",
    )
    h(doc, "8.1 场景建议矩阵", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[23]])
    h(doc, "8.2 制作方式", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[24]])
    h(doc, "8.3 交付项", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[25]])
    h(doc, "8.4 验收标准", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[26]])

    # —— 业务流程 ——
    h(doc, "9 核心业务流程（逐步清单）", 1)
    sections = [
        ("9.1 用户注册登录", 134, 143),
        ("9.2 买家购买", 146, 158),
        ("9.3 卖家发布", 160, 170),
        ("9.4 砍价与出价", 172, 177),
        ("9.5 合同签署", 179, 183),
        ("9.6 提现", 185, 191),
        ("9.7 人工估号", 195, 199),
    ]
    for title, start_i, end_i in sections:
        h(doc, title, 2)
        for k in range(start_i, end_i + 1):
            if k not in pmap:
                continue
            text = pmap[k].strip()
            if not text:
                continue
            if text.startswith("约束规则：") or text.startswith("动效边界："):
                p(doc, text)
                continue
            if text.startswith("-"):
                doc.add_paragraph(text, style="List Bullet")
            elif re.match(r"^\d+\.", text):
                doc.add_paragraph(text, style="List Number")
            else:
                p(doc, text)
    h(doc, "9.8 代理授权", 2)
    agent_steps = []
    for k in range(201, 209):
        if k in pmap:
            line = pmap[k].strip()
            if line:
                agent_steps.append(line)
    for n, line in enumerate(agent_steps, 1):
        doc.add_paragraph(f"{n}. {line}", style="List Number")

    # —— 功能需求 ——
    h(doc, "10 功能需求（按子域拆分，整合稿表格原样归并）", 1)
    p(
        doc,
        "字段级与界面细节以签字设计稿及后台属性模板为准；下列表格按整合稿章节拆分，优先保证可追溯。",
    )
    h(doc, "10.1 UI 设计与前端切图", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[27]])
    h(doc, "10.2 登录注册", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[28]])
    h(doc, "10.3 首页", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[29]])
    h(doc, "10.4 买点什么", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[30]])
    h(doc, "10.5 商品筛选与排序", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[31]])
    h(doc, "10.6 商品详情页", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[32]])
    h(doc, "10.7 下单支付与订单详情", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[33]])
    h(doc, "10.8 卖点什么 / 发布", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[34]])
    h(doc, "10.9 我的与个人资产", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[35]])
    h(doc, "10.10 积分商城", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[36]])
    h(doc, "10.11 推广邀请", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[37]])
    h(doc, "10.12 支付与分账", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[38]])
    h(doc, "10.13 消息与 IM", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[39]])
    h(doc, "10.14 我的订单", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[40]])
    h(doc, "10.15 设置与其他", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[41]])
    h(doc, "10.16 管理后台模块", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[42]])

    # —— 状态机 ——
    h(doc, "11 状态机摘要", 1)
    h(doc, "11.1 商品", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[43]])
    h(doc, "11.2 订单", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[44]])
    h(doc, "11.3 提现", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[45]])
    h(doc, "11.4 售后", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[46]])

    # —— 数据模型 ——
    h(doc, "12 逻辑数据对象与字段（初稿）", 1)
    p(doc, "字段命名沿用整合稿英文标识，实施阶段可随数据库设计微调；含义以业务域评审为准。")
    for title, idx in [
        ("12.1 用户 User", 47),
        ("12.2 游戏 Game", 48),
        ("12.3 商品 Product", 49),
        ("12.4 订单 Order", 50),
        ("12.5 钱包 WalletTransaction", 51),
        ("12.6 优惠券 Coupon", 52),
        ("12.7 推广 InviteReward", 53),
    ]:
        h(doc, title, 2)
        table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[idx]])

    # —— API 职责摘要 ——
    h(doc, "13 服务端能力清单（职责摘要）", 1)
    p(
        doc,
        "整合稿曾给出示例 URL；出于版本漂移与评审噪声考虑，V2 仅保留职责描述。路径、鉴权态、幂等等在《接口说明》单独立项。",
    )
    domains = [
        "用户与认证",
        "游戏与商品",
        "订单与议价",
        "钱包与优惠",
        "消息与 IM",
        "管理后台",
    ]
    table_grid(doc, api_duty_table(T, domains, 54, 60))

    # —— 第三方 ——
    h(doc, "14 第三方服务与外部依赖", 1)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[60]])

    # —— 非功能 ——
    h(doc, "15 非功能需求", 1)
    h(doc, "15.1 性能", 2)
    bullets(
        doc,
        [
            strip_bullet(pmap[i])
            for i in range(250, 255)
            if i in pmap and pmap[i].strip().startswith("-")
        ],
    )
    h(doc, "15.2 安全", 2)
    bullets(
        doc,
        [strip_bullet(pmap[i]) for i in range(256, 263) if i in pmap and pmap[i].strip().startswith("-")],
    )
    h(doc, "15.3 合规", 2)
    bullets(
        doc,
        [strip_bullet(pmap[i]) for i in range(265, 270) if i in pmap and pmap[i].strip().startswith("-")],
    )
    h(doc, "15.4 兼容矩阵（整合稿）", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[61]])

    # —— 验收 ——
    h(doc, "16 验收标准", 1)
    h(doc, "16.1 UI 与前端", 2)
    bullets(
        doc,
        [strip_bullet(pmap[i]) for i in range(273, 280) if i in pmap and pmap[i].strip().startswith("-")],
    )
    h(doc, "16.2 业务功能", 2)
    bullets(
        doc,
        [strip_bullet(pmap[i]) for i in range(281, 291) if i in pmap and pmap[i].strip().startswith("-")],
    )
    h(doc, "16.3 管理后台", 2)
    bullets(
        doc,
        [strip_bullet(pmap[i]) for i in range(292, 298) if i in pmap and pmap[i].strip().startswith("-")],
    )

    # —— 阶段计划 ——
    h(doc, "17 阶段计划（整合稿工作分解）", 1)
    h(doc, "17.1 阶段一：需求与 UI 设计", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[62]])
    h(doc, "17.2 阶段二：前端框架与页面实现", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[63]])
    h(doc, "17.3 阶段三：后端与后台", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[64]])
    h(doc, "17.4 阶段四：联调、测试与上线", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[65]])

    # —— 页面清单 ——
    h(doc, "18 页面清单初稿", 1)
    h(doc, "18.1 前台用户端", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[66]])
    h(doc, "18.2 管理后台", 2)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[67]])

    # —— 风险 ——
    h(doc, "19 风险与注意事项", 1)
    table_grid(doc, [[c.replace("\n", " ") for c in row] for row in T[68]])

    # —— 待定 ——
    h(doc, "20 待定问题清单（评审输出）", 1)
    bullets(
        doc,
        [
            "“中介担保”交易类型的前台流程、后台字段与资金清算规则（当前为预留入口）。",
            "四方支付服务商、费率、分账角色字典与对账口径。",
            "IM 选型、消息留存周期、审计与水印策略。",
            "电子签模板字段、签署顺序、主体信息与法务定稿。",
            "包赔套餐的赔付上限、免责声明与展示文案。",
            "代理毁约计数默认阈值（整合稿示例为 5 次）及扣款比例（示例 30% ，可配置）是否采纳。",
            "人脸识别服务商及活体策略。",
            "管理后台是否强制暗黑皮肤或允许高效浅色主题（效率向）。",
            "整合稿“测试部署”线是否包含驻场运维与免费维保周期：商务另行确认，不在 PRD 正文收口。",
        ],
    )

    h(doc, "附录 A 报价与商务", 1)
    p(
        doc,
        "本书不包含金额口径。商务报价、人天单价、税费与付款里程碑以独立《报价单/合同附件》为准，本附录仅作交叉引用索引。",
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
